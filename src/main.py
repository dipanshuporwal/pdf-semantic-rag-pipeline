import os
import time
from typing import List, Tuple
from dotenv import load_dotenv
import pdfplumber
import pickle

from langchain_groq import ChatGroq
from langchain_experimental.text_splitter import SemanticChunker
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document as LangDocument

from pymilvus import (
    connections,
    Collection,
    CollectionSchema,
    FieldSchema,
    DataType,
    utility,
)

from rank_bm25 import BM25Okapi
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document

# Load environment
load_dotenv()

# Cache files
CHUNK_FILE = "cached_chunks.pkl"
EMBED_FLAG = "embeddings_inserted.flag"


# ----- Step 1: Extract text from PDFs -----
def extract_text_from_pdfs(pdf_dir: str) -> List[str]:
    texts = []
    for file in os.listdir(pdf_dir):
        if file.endswith(".pdf"):
            with pdfplumber.open(os.path.join(pdf_dir, file)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        texts.append(text)
    return texts


# ----- Cache Load/Save for chunks -----
def save_chunks(chunks: List[LangDocument]):
    with open(CHUNK_FILE, "wb") as f:
        pickle.dump(chunks, f)


def load_chunks() -> List[LangDocument]:
    if os.path.exists(CHUNK_FILE):
        with open(CHUNK_FILE, "rb") as f:
            return pickle.load(f)
    return []


# ----- Step 2: Create HuggingFace Embedder -----
def get_embedder() -> HuggingFaceEmbeddings:
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-mpnet-base-v2"
    )


# ----- Step 3: Chunk texts with caching -----
def semantic_chunking(texts: List[str], embedder) -> List[LangDocument]:
    cached = load_chunks()
    if cached:
        print("Loaded chunks from cache.")
        return cached
    print("Chunking texts...")
    splitter = SemanticChunker(embedder)
    chunks = splitter.create_documents(texts)
    save_chunks(chunks)
    return chunks


# ----- Step 4: Connect to Milvus -----
def connect_milvus():
    connections.connect("default", host="localhost", port="19530")


# ----- Step 5: Create or load Milvus Collection -----
def create_collection(name: str, index_type: str) -> Collection:
    # if utility.has_collection(name):
    #     utility.drop_collection(name)
    if utility.has_collection(name):
        print(f"Collection '{name}' already exists. Loading...")
        collection = Collection(name)
        collection.load()
        return collection

    print(f"Creating collection '{name}'...")
    schema = CollectionSchema(
        [
            FieldSchema(
                name="id", dtype=DataType.INT64, is_primary=True, auto_id=True
            ),
            FieldSchema(
                name="embedding", dtype=DataType.FLOAT_VECTOR, dim=768
            ),
            FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=65535),
        ]
    )
    collection = Collection(name=name, schema=schema)

    index_params = {
        "FLAT": {"index_type": "FLAT", "metric_type": "L2", "params": {}},
        "IVF_FLAT": {
            "index_type": "IVF_FLAT",
            "metric_type": "L2",
            "params": {"nlist": 128},
        },
        "HNSW": {
            "index_type": "HNSW",
            "metric_type": "L2",
            "params": {"M": 16, "efConstruction": 200},
        },
    }[index_type]

    collection.create_index("embedding", index_params)
    collection.load()
    return collection


# ----- Step 6: Insert Embeddings with flag check -----
def insert_embeddings(
    collection: Collection, embeddings: List[List[float]], texts: List[str]
):
    if os.path.exists(EMBED_FLAG):
        print("Embeddings already inserted. Skipping insertion.")
        return
    print("Inserting embeddings into Milvus...")
    collection.insert([embeddings, texts])
    with open(EMBED_FLAG, "w") as f:
        f.write("done")


# ----- Step 7: Search Milvus Collection -----
def search_collection(
    collection: Collection, query_vector: List[float], k: int = 5
):
    results = collection.search(
        [query_vector],
        "embedding",
        param={"metric_type": "L2"},
        limit=k,
        output_fields=["text"],
    )
    return results


# ----- Step 8: BM25 Reranking -----
def bm25_rerank(query: str, docs: List[str]) -> List[str]:
    tokenized = [doc.lower().split() for doc in docs]
    bm25 = BM25Okapi(tokenized)
    scores = bm25.get_scores(query.lower().split())
    ranked = [doc for _, doc in sorted(zip(scores, docs), reverse=True)]
    return ranked


# ----- Step 9: MMR Reranking -----
def mmr(
    query_embedding: List[float],
    doc_embeddings: List[List[float]],
    docs: List[str],
    lambda_param=0.5,
    k=5,
) -> List[str]:
    selected_idxs, candidate_idxs = [], list(range(len(doc_embeddings)))
    while candidate_idxs and len(selected_idxs) < k:
        best_idx, best_score = -1, -float("inf")
        for idx in candidate_idxs:
            relevance = cosine_similarity(
                [query_embedding], [doc_embeddings[idx]]
            )[0][0]
            diversity = max(
                [
                    cosine_similarity(
                        [doc_embeddings[idx]], [doc_embeddings[sel]]
                    )[0][0]
                    for sel in selected_idxs
                ]
                or [0]
            )
            score = lambda_param * relevance - (1 - lambda_param) * diversity
            if score > best_score:
                best_idx, best_score = idx, score
        selected_idxs.append(best_idx)
        candidate_idxs.remove(best_idx)
        print(selected_idxs)
        print(len(docs))
    return [docs[i] for i in selected_idxs]


# ----- Step 10: Ask LLM -----
def ask_llm(query: str, context: str) -> str:
    llm = ChatGroq(temperature=0.2, model_name="llama-3.3-70b-versatile")
    prompt = (
        "Answer the question based on the following"
        f" context:\n\n{context}\n\nQuestion: {query}"
    )
    return llm.invoke(prompt).content


# ----- Step 11: Save to DOCX -----
def write_to_docx(answer: str, file_path: str):
    doc = Document()
    doc.add_heading("LLM Response", 0)
    doc.add_paragraph(answer)
    doc.save(file_path)


# ----- Main Pipeline -----
def full_pipeline(
    pdf_dir: str,
    user_query: str,
    index_type="HNSW",
    collection_name="pdf_embeddings",
    docx_output="response.docx",
):
    print("🔍 Extracting text...")
    texts = extract_text_from_pdfs(pdf_dir)

    print("🔗 Embedding and Chunking...")
    embedder = get_embedder()
    docs = semantic_chunking(texts, embedder)
    doc_texts = [doc.page_content for doc in docs]
    doc_embeddings = embedder.embed_documents(doc_texts)

    print("🧠 Setting up Milvus...")
    connect_milvus()
    collection = create_collection(collection_name, index_type)
    insert_embeddings(collection, doc_embeddings, doc_texts)

    print("🧾 Embedding query...")
    query_embedding = embedder.embed_query(user_query)

    print("🔎 Retrieving from Milvus...")
    results = search_collection(collection, query_embedding)
    print(results)
    # result.id is Milvus internal ID, not the doc index, so use result.entity.get("id") or result.primary_keys if needed.
    # Here, since we have auto_id in Milvus, the result.id corresponds to entity id which is auto_id but
    # doc_texts is a list indexed from 0, so you can't use result.id as index directly.
    # So let's fetch top-k by order from results and map to docs in order:
    print("🔎 Retrieving from Milvus...")
    hits = results[0]  # Get the first query's search results
    top_texts = [
        hit.entity.get("text") for hit in hits
    ]  # hit.id should work here
    print(top_texts)

    print("🔁 Reranking with BM25 & MMR...")
    bm25_ranked = bm25_rerank(user_query, top_texts)
    print(f"bm25_ranked: {bm25_ranked}")
    # mmr_ranked = mmr(query_embedding, doc_embeddings, bm25_ranked)
    # print(f"mmr_ranked: {mmr_ranked}")

    context = "\n\n".join(bm25_ranked)
    print("🤖 Asking LLM...")
    answer = ask_llm(user_query, context)

    print("💾 Writing to DOCX...")
    write_to_docx(answer, docx_output)

    print("✅ Done! Answer saved to", docx_output)


if __name__ == "__main__":
    full_pipeline(
        pdf_dir="src/data",  # Folder with PDFs
        user_query="What is reinforcement learning?",
        index_type="HNSW",  # Options: FLAT, IVF_FLAT, HNSW
        docx_output="ISO_Compliance_Response.docx",
    )
